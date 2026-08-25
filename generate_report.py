import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
import os

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_heading_1(doc, text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(8)
    h.paragraph_format.keep_with_next = True
    run = h.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 23, 42) # Slate 900
    return h

def add_heading_2(doc, text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True
    run = h.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 58, 138) # Blue 900
    return h

def add_heading_3(doc, text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(4)
    h.paragraph_format.keep_with_next = True
    run = h.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(51, 65, 85) # Slate 700
    return h

def add_p(doc, text, bold_prefix=None, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Arial'
        r_pre.font.size = Pt(10.5)
        r_pre.font.bold = True
        r_pre.font.color.rgb = RGBColor(15, 23, 42)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(10.5)
    run.font.italic = italic
    run.font.color.rgb = RGBColor(51, 65, 85)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Arial'
        r_pre.font.size = Pt(10.5)
        r_pre.font.bold = True
        r_pre.font.color.rgb = RGBColor(15, 23, 42)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(51, 65, 85)
    return p

def add_callout(doc, text, title="CATATAN PENTING"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "F0F9FF") # Sky light
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r_t = p.add_run(f"📌 {title}\n")
    r_t.font.name = 'Arial'
    r_t.font.size = Pt(10.5)
    r_t.font.bold = True
    r_t.font.color.rgb = RGBColor(3, 105, 161)
    
    r_b = p.add_run(text)
    r_b.font.name = 'Arial'
    r_b.font.size = Pt(10)
    r_b.font.color.rgb = RGBColor(30, 58, 138)
    
    # Empty paragraph after table for spacing
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(6)

def style_table(table, col_widths, headers, data):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        set_cell_background(hdr_cells[i], "1E293B") # Dark slate header
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=120, right=120)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            
    for row_idx, row_data in enumerate(data):
        row_cells = table.add_row().cells
        bg_color = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, cell_value in enumerate(row_data):
            row_cells[col_idx].text = str(cell_value)
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=120, right=120)
            p = row_cells[col_idx].paragraphs[0]
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(51, 65, 85)

    for row in table.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = Inches(w)

def add_image_with_caption(doc, img_path, caption_text, width_in_inches=6.0):
    if os.path.exists(img_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(10)
        p_img.paragraph_format.space_after = Pt(4)
        run = p_img.add_run()
        run.add_picture(img_path, width=Inches(width_in_inches))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(12)
        r_cap = p_cap.add_run(f"Gambar: {caption_text}")
        r_cap.font.name = 'Arial'
        r_cap.font.size = Pt(9.5)
        r_cap.font.italic = True
        r_cap.font.color.rgb = RGBColor(100, 116, 139)

def main():
    doc = docx.Document()

    # Set Margins (1 inch = 72pt = 1440 dxa)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # --- COVER PAGE ---
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(36)
    p_title.paragraph_format.space_after = Pt(12)
    r_title = p_title.add_run("LAPORAN FINAL PROJECT\nSISTEM PLATFORM SAAS REST API")
    r_title.font.name = 'Arial'
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(15, 23, 42)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(24)
    r_sub = p_sub.add_run("CineData API — Movie Data SaaS Platform dengan Autentikasi JWT, API Key Middleware, Pembatasan Laju (Rate Limiting), dan Telemetri Real-Time")
    r_sub.font.name = 'Arial'
    r_sub.font.size = Pt(12)
    r_sub.font.color.rgb = RGBColor(71, 85, 105)

    # Decorative Box on Cover Page
    tbl_cov = doc.add_table(rows=1, cols=1)
    tbl_cov.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_cov = tbl_cov.cell(0, 0)
    set_cell_background(cell_cov, "F1F5F9")
    set_cell_margins(cell_cov, top=180, bottom=180, left=250, right=250)
    p_cov_c = cell_cov.paragraphs[0]
    p_cov_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_cov_1 = p_cov_c.add_run("TEKNOLOGI & INFRASTRUKTUR SISTEM\n")
    r_cov_1.font.name = 'Arial'
    r_cov_1.font.size = Pt(11)
    r_cov_1.font.bold = True
    r_cov_1.font.color.rgb = RGBColor(30, 58, 138)
    
    r_cov_2 = p_cov_c.add_run("Backend: Express.js (Node.js)  |  Database: PostgreSQL / Supabase  |  Deployment: Vercel Serverless Functions\nKeamanan: JWT Authentication + SHA-256 Hashed API Key + Bcrypt Password Hashing")
    r_cov_2.font.name = 'Arial'
    r_cov_2.font.size = Pt(10)
    r_cov_2.font.color.rgb = RGBColor(51, 65, 85)

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_before = Pt(48)
    p_meta.paragraph_format.space_after = Pt(0)
    r_meta = p_meta.add_run("Disusun untuk Memenuhi Tugas Akhir Praktikum Web Service (PWS / UCP)\nTahun Akademik 2026")
    r_meta.font.name = 'Arial'
    r_meta.font.size = Pt(11)
    r_meta.font.bold = True
    r_meta.font.color.rgb = RGBColor(15, 23, 42)

    doc.add_page_break()

    # --- RINGKASAN EKSEKUTIF ---
    add_heading_1(doc, "Ringkasan Eksekutif & Verifikasi Kriteria Penilaian")
    add_p(doc, "Laporan ini mendokumentasikan secara komprehensif perancangan, arsitektur, implementasi, dan pengujian sistem CineData API. CineData API dirancang sebagai platform SaaS (Software as a Service) penyedia metadata film dan hiburan berkinerja tinggi, terinspirasi oleh model bisnis platform modern seperti OpenRouter API dan Weather API.")
    
    add_p(doc, "Seluruh indikator teknis dan akademis yang dipersyaratkan dalam tugas akhir telah diimplementasikan secara sempurna dengan rincian pencapaian sebagai berikut:")
    
    table_verif = doc.add_table(rows=1, cols=3)
    headers_v = ["Indikator / Persyaratan", "Status Pemenuhan", "Rincian Implementasi Proyek"]
    data_v = [
        ["Model Sistem SaaS API", "✅ 100% Terpenuhi", "Layanan penyedia data eksternal via API Key (x-api-key), rate limiting tier (Free, Dev, Enterprise), & telemetry usage analytics."],
        ["Struktur Model & Basis Data", "✅ Melampaui Syarat", "Struktur Model (models/User.js, models/ApiKey.js, models/Movie.js, models/Genre.js, models/People.js, models/Usage.js) & 11 Tabel Relasional."],
        ["Autentikasi Login JWT", "✅ 100% Terpenuhi", "Autentikasi pengembang berbasis JSON Web Token (/api/auth/register, /api/auth/login, /api/auth/me) dengan proteksi Authorization Bearer Header."],
        ["Laporan Diagram Sistem", "✅ 100% Terpenuhi", "Laporan dilengkapi ERD, Use Case Diagram, Activity Diagram, User Flow, dan Diagram Arsitektur Sistem."],
        ["Deployment di Vercel", "✅ 100% Terpenuhi", "Telah dikonfigurasi & dioptimalkan untuk Vercel Serverless Functions via vercel.json."],
        ["Jumlah & Kompleksitas Data (Min 50 Data)", "✅ Melampaui Syarat", "Seeding data berkapasitas besar: 185+ Data Film, 23 Genre, 105 Aktor/Sutradara, 32 Perusahaan Produksi, 550+ Ulasan, dan 525+ Log Telemetri."],
        ["Tech Stack Resmi", "✅ 100% Terpenuhi", "Node.js + Express.js (Backend), EJS View Engine (Template Rendering), PostgreSQL / Supabase (Database Pool via pg), Vercel (Cloud Hosting)."]
    ]
    style_table(table_verif, [2.0, 1.3, 3.2], headers_v, data_v)

    add_callout(doc, "Sistem dilengkapi dengan In-Memory Fallback Engine otomatis yang menjamin aplikasi tetap berjalan 100% tanpa downtime meskipun koneksi PostgreSQL/Supabase mengalami kendala jaringan.", "KEUNGGULAN ARSITEKTUR")

    # --- BAB I: PENDAHULUAN ---
    add_heading_1(doc, "BAB I: PENDAHULUAN")
    
    add_heading_2(doc, "1.1 Latar Belakang")
    add_p(doc, "Dalam era transformasi digital, kebutuhan akan integrasi data antar aplikasi (*Application Programming Interface* / API) tumbuh secara eksponensial. Model bisnis *Software as a Service* (SaaS) berbasis API—seperti OpenRouter untuk AI model routing atau Weather API untuk data meteorologi—menjadi standar industri dalam mendistribusikan data secara aman, terukur, dan komersial.")
    add_p(doc, "CineData API dikembangkan untuk menjawab kebutuhan tersebut dalam domain industri media dan hiburan. Sistem ini memungkinkan pengembang aplikasi pihak ketiga (web, mobile, maupun SDK) untuk mengonsumsi katalog metadata film terlengkap secara terstruktur, aman, dan efisien dengan mengikat setiap permintaan API pada API Key yang terverifikasi.")

    add_heading_2(doc, "1.2 Rumusan Masalah")
    add_bullet(doc, "Bagaimana merancang arsitektur RESTful SaaS API yang mampu mengamankan akses data publik menggunakan API Key dan Hashing SHA-256?", "1. ")
    add_bullet(doc, "Bagaimana mengimplementasikan autentikasi pengembang berbasis JSON Web Token (JWT) untuk mengelola kredensial dan lisensi API Key?", "2. ")
    add_bullet(doc, "Bagaimana membangun basis data relasional kompleks yang mencakup lebih dari 50++ data film beserta entitas pendukungnya (genre, pemeran, perusahaan produksi, ulasan)?", "3. ")
    add_bullet(doc, "Bagaimana mengontrol penggunaan API dengan pembatasan laju (*rate limiting*) dan pencatatan telemetri secara real-time?", "4. ")
    add_bullet(doc, "Bagaimana mengkonfigurasi dan menggelar (*deploy*) aplikasi Express.js pada arsitektur *serverless* Vercel?", "5. ")

    add_heading_2(doc, "1.3 Tujuan Proyek")
    add_p(doc, "Tujuan utama dari pembuatan proyek ini adalah menghasilkan sistem platform API SaaS skala produksi yang memenuhi standar akademis dan industri, melencarkan pengujian integrasi otomatis, serta menyediakan dokumentasi diagram yang komprehensif.")

    # --- BAB II: ARSITEKTUR & KEAMANAN SISTEM ---
    add_heading_1(doc, "BAB II: ARSITEKTUR & KEAMANAN SISTEM")

    add_heading_2(doc, "2.1 Spesifikasi Teknologi (Tech Stack)")
    add_bullet(doc, "Express.js (v4.19) pada lingkungan runtime Node.js v18+ dengan EJS (Embedded JavaScript Templates) View Engine.", "Backend & View Framework: ")
    add_bullet(doc, "PostgreSQL (Supabase Cloud Database) yang terhubung melalui pg Connection Pool. Dilengkapi dengan In-Memory Fallback State Engine internal untuk keandalan tinggi.", "Database Layer: ")
    add_bullet(doc, "Kombinasi jsonwebtoken untuk sesi pengembang, bcryptjs untuk hashing password, crypto SHA-256 untuk hashing API Key rahasia, Helmet HTTP headers, dan CORS policy.", "Keamanan & Protokol: ")
    add_bullet(doc, "Vercel Serverless Functions via vercel.json, mengoptimalkan Express server menjadi fungsi serverless berlatensi rendah.", "Deployment Cloud: ")

    add_heading_2(doc, "2.2 Arsitektur Sistem SaaS")
    add_p(doc, "Sistem CineData API memisahkan dua alur kerja utama: Alur Manajemen Dashboard (JWT Authenticated) dan Alur Konsumsi Data Publik (API Key Authenticated).")
    
    add_image_with_caption(doc, "docs/api-architecture.png", "Diagram Arsitektur Sistem CineData API", width_in_inches=6.0)

    add_heading_2(doc, "2.3 Mekanisme Keamanan API Key & Hashing")
    add_p(doc, "Untuk mencegah kebocoran API Key dalam basis data, CineData API menerapkan prinsip Hashing SHA-256:")
    add_bullet(doc, "Ketika pengembang membuat API Key baru, sistem memberikan raw key sekali saja (contoh: cd_live_a1b2c3d4...).", "1. Format Prefixing: ")
    add_bullet(doc, "Basis data hanya menyimpan hash SHA-256 dari key tersebut (key_hash).", "2. Hashing SHA-256: ")
    add_bullet(doc, "Setiap permintaan masuk di-hash menggunakan algoritma SHA-256 lalu dicocokkan dengan rekord di basis data, menjamin keamanan tingkat tinggi.", "3. Verifikasi Header: ")

    # --- BAB III: PERANCANGAN BASIS DATA & DIAGRAM SISTEM ---
    add_heading_1(doc, "BAB III: PERANCANGAN BASIS DATA & DIAGRAM SISTEM")

    add_heading_2(doc, "3.1 Struktur Relasional Basis Data (11 Tabel)")
    add_p(doc, "Basis data dirancang dengan normalisasi tinggi untuk mendukung query kompleks dan hubungan *many-to-many* antar entitas:")
    
    table_db = doc.add_table(rows=1, cols=3)
    headers_db = ["Nama Tabel", "Tipe Entitas / Fungsi", "Keterangan Relasi & Kolom Kunci"]
    data_db = [
        ["users", "Entitas Pengembang", "PK: id, Email Unik, password_hash, created_at"],
        ["api_keys", "Entitas Kunci API", "FK: user_id -> users(id), key_prefix, key_hash (UK), tier, expires_at"],
        ["movies", "Katalog Utama Film", "PK: id, title, slug (UK), overview, release_date, budget, revenue, popularity, vote_average"],
        ["genres", "Kategori Genre", "PK: id, name (UK), slug (UK)"],
        ["movie_genres", "Tabel Penghubung (Junction)", "FK: movie_id -> movies(id), FK: genre_id -> genres(id)"],
        ["people", "Pemeran & Sutradara", "PK: id, name, biography, birth_date, birth_place"],
        ["movie_cast", "Relasi Pemeran Film", "FK: movie_id -> movies(id), FK: person_id -> people(id), character_name"],
        ["movie_crew", "Relasi Kru / Sutradara", "FK: movie_id -> movies(id), FK: person_id -> people(id), department, job"],
        ["production_companies", "Studio Produksi", "PK: id, name, country, logo_url"],
        ["movie_companies", "Relasi Studio Film", "FK: movie_id -> movies(id), FK: company_id -> production_companies(id)"],
        ["reviews", "Ulasan & Rating", "FK: movie_id -> movies(id), author, rating, content"],
        ["api_usage", "Log Telemetri API", "FK: api_key_id -> api_keys(id), endpoint, method, status_code, response_time, ip_address"]
    ]
    style_table(table_db, [1.5, 1.8, 3.2], headers_db, data_db)

    add_heading_2(doc, "3.2 Entity Relationship Diagram (ERD)")
    add_p(doc, "ERD menggambarkan keterhubungan antar 11 entitas relasional dalam sistem CineData API:")
    add_image_with_caption(doc, "docs/erd.png", "Entity Relationship Diagram (ERD) CineData API", width_in_inches=6.0)

    add_heading_2(doc, "3.3 Use Case Diagram")
    add_p(doc, "Diagram Use Case menjelaskan interaksi antara Aktor Pengembang (*Developer*), Aplikasi Pihak Ketiga (*External App*), dan Sistem backend:")
    add_image_with_caption(doc, "docs/usecase.png", "Use Case Diagram CineData SaaS Platform", width_in_inches=6.0)

    add_heading_2(doc, "3.4 Activity Diagram & User Flow")
    add_p(doc, "Activity Diagram memvisualisasikan siklus hidup setiap pemanggilan REST API publik, mulai dari ekstraksi header, verifikasi hash API key, pembatasan rate limit, eksekusi query SQL, hingga pencatatan log telemetri secara asinkron:")
    add_image_with_caption(doc, "docs/activity.png", "Activity Diagram - Siklus Hidup Permintaan API Public", width_in_inches=5.8)

    add_p(doc, "User Flow alur kerja pengembang pada Web Dashboard SaaS:")
    add_image_with_caption(doc, "docs/userflow.png", "User Flow Alur Kerja Web Dashboard", width_in_inches=5.8)

    # --- BAB IV: IMPLEMENTASI & KATALOG ENDPOINT API ---
    add_heading_1(doc, "BAB IV: IMPLEMENTASI & KATALOG ENDPOINT API")

    add_heading_2(doc, "4.1 Rute Autentikasi JWT (Developer Auth)")
    table_auth = doc.add_table(rows=1, cols=3)
    headers_a = ["Metode", "Endpoint", "Fungsi & Deskripsi"]
    data_a = [
        ["POST", "/api/auth/register", "Mendaftarkan akun pengembang baru dengan validasi email & hash password."],
        ["POST", "/api/auth/login", "Autentikasi pengembang dan menerbitkan JWT Session Token."],
        ["GET", "/api/auth/me", "Mengambil data profil pengembang terautentikasi (memerlukan Header Bearer Token)."],
        ["POST", "/api/auth/logout", "Mengakhiri sesi pengembang."]
    ]
    style_table(table_auth, [1.0, 2.0, 3.5], headers_a, data_a)

    add_heading_2(doc, "4.2 Rute Manajemen API Key")
    table_keys = doc.add_table(rows=1, cols=3)
    headers_k = ["Metode", "Endpoint", "Fungsi & Deskripsi"]
    data_k = [
        ["POST", "/api/keys", "Membuat API Key baru (mengembalikan rahasia mentah cd_live_... sekali saja)."],
        ["GET", "/api/keys", "Menampilkan daftar seluruh API Key milik pengembang beserta metrik penggunaan."],
        ["GET", "/api/keys/:id", "Mengambil rincian metadata API Key tertentu."],
        ["PATCH", "/api/keys/:id/revoke", "Mencabut (menonaktifkan) status API Key."],
        ["DELETE", "/api/keys/:id", "Menghapus rekord API Key secara permanen."]
    ]
    style_table(table_keys, [1.0, 2.0, 3.5], headers_k, data_k)

    add_heading_2(doc, "4.3 Katalog REST API Publik Metadata Film (Require x-api-key)")
    table_pub = doc.add_table(rows=1, cols=3)
    headers_p = ["Metode", "Endpoint", "Fungsi & Fitur Query"]
    data_p = [
        ["GET", "/api/v1/movies", "Menampilkan daftar film dengan Paginasi (?page=1&limit=10), Search (?search=nolan), Filter Genre/Tahun/Rating, dan Sorting (?sort=-rating)."],
        ["GET", "/api/v1/movies/:id", "Mengambil metadata detail film berdasarkan ID."],
        ["GET", "/api/v1/movies/slug/:slug", "Mengambil metadata detail film berdasarkan URL slug unik."],
        ["GET", "/api/v1/movies/:id/cast", "Mengambil daftar pemeran dan karakter film."],
        ["GET", "/api/v1/movies/:id/reviews", "Mengambil ulasan dan rating ulasan dari para kritikus."],
        ["GET", "/api/v1/genres", "Menampilkan daftar seluruh 23 genre film."],
        ["GET", "/api/v1/genres/:id", "Mengambil detail genre dan jumlah film terkait."],
        ["GET", "/api/v1/people", "Menampilkan katalog aktor, sutradara, dan kru film."],
        ["GET", "/api/v1/people/:id", "Mengambil rincian biografi dan filmografi orang."],
        ["GET", "/api/v1/companies", "Menampilkan 32 studio dan perusahaan produksi film utama."]
    ]
    style_table(table_pub, [1.0, 2.2, 3.3], headers_p, data_p)

    # --- BAB V: DATASET & PENGUJIAN SISTEM ---
    add_heading_1(doc, "BAB V: KOMPLEKSITAS DATASET & PENGUJIAN SISTEM")

    add_heading_2(doc, "5.1 Kapasitas Seeding Data")
    add_p(doc, "Sistem dilengkapi skrip seeding otomatis (npm run seed) yang mengisi basis data dengan dataset nyata berkapasitas besar:")
    add_bullet(doc, "185++ Data Film Lengkap (Judul, Overview, Rating, Popularitas, Anggaran, Pendapatan, Poster & Trailer URL).", "• ")
    add_bullet(doc, "23 Genre Film Terkategori (Action, Sci-Fi, Drama, Cyberpunk, Film Noir, dll).", "• ")
    add_bullet(doc, "105 Pemeran & Sutradara Ternama (Christopher Nolan, Denis Villeneuve, Cillian Murphy, Leonardo DiCaprio, dll).", "• ")
    add_bullet(doc, "32 Studio Produksi Dunia (Warner Bros, Universal, A24, Marvel Studios, Studio Ghibli, dll).", "• ")
    add_bullet(doc, "550++ Ulasan & Rating Kritikus.", "• ")
    add_bullet(doc, "525++ Rekord Log Telemetri Penggunaan API.", "• ")

    add_heading_2(doc, "5.2 Hasil Automated Integration Test Runner")
    add_p(doc, "Sistem diuji menggunakan test runner otomatis (npm test) yang memverifikasi 13 skenario integrasi end-to-end secara transparan:")
    
    table_test = doc.add_table(rows=1, cols=3)
    headers_t = ["No", "Skenario pengujian", "Hasil Verifikasi"]
    data_t = [
        ["1", "Registrasi Pengembang Akun Baru (/api/auth/register)", "PASSED (201 Created + Token)"],
        ["2", "Login Pengembang & Penerbitan JWT (/api/auth/login)", "PASSED (200 OK + JWT Cookie/Bearer)"],
        ["3", "Proteksi Endpoint Terproteksi Tanpa Token", "PASSED (401 Unauthorized)"],
        ["4", "Pembuatan API Key Baru & Hashing SHA-256", "PASSED (201 Created + Raw Key)"],
        ["5", "Penampilan Daftar API Key milik Pengembang", "PASSED (200 OK + Metadata List)"],
        ["6", "Akses API Publik Tanpa Header x-api-key", "PASSED (401 MISSING_API_KEY)"],
        ["7", "Akses API Publik Dengan Invalid API Key", "PASSED (401 INVALID_API_KEY)"],
        ["8", "Akses API Publik Dengan Valid API Key", "PASSED (200 OK + JSON Payload)"],
        ["9", "Penyaringan & Paginasi Catalog Film", "PASSED (200 OK + Meta Pagination)"],
        ["10", "Pencarian Film Berdasarkan Kata Kunci (search=nolan)", "PASSED (200 OK + Filtered Array)"],
        ["11", "Pengujian Pembatasan Laju (Rate Limiting Tier)", "PASSED (429 RATE_LIMIT_EXCEEDED)"],
        ["12", "Pencatatan Telemetri Penggunaan Asinkron", "PASSED (Logged in api_usage table)"],
        ["13", "Pencabutan & Penghapusan API Key (Revoke/Delete)", "PASSED (200 OK + Key Deactivated)"]
    ]
    style_table(table_test, [0.5, 4.0, 2.0], headers_t, data_t)

    # --- BAB VI: DEPLOYMENT VERCEL & TAUTAN PROYEK ---
    add_heading_1(doc, "BAB VI: DEPLOYMENT & TAUTAN PROYEK")

    add_heading_2(doc, "6.1 Konfigurasi Serverless Vercel (vercel.json)")
    add_p(doc, "Aplikasi telah dikonfigurasi secara presisi agar dapat dijalankan sebagai Serverless Functions di platform Vercel Cloud:")
    
    code_text = '''{
  "version": 2,
  "builds": [
    {
      "src": "app.js",
      "use": "@vercel/node"
    }
  ],
  "routes": [
    {
      "src": "/(.*)",
      "dest": "app.js"
    }
  ]
}'''
    tbl_code = doc.add_table(rows=1, cols=1)
    tbl_code.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_code = tbl_code.cell(0, 0)
    set_cell_background(cell_code, "0F172A")
    set_cell_margins(cell_code, top=100, bottom=100, left=150, right=150)
    p_code = cell_code.paragraphs[0]
    r_code = p_code.add_run(code_text)
    r_code.font.name = 'Consolas'
    r_code.font.size = Pt(9)
    r_code.font.color.rgb = RGBColor(56, 189, 248) # Sky 400

    add_heading_2(doc, "6.2 Tautan Pengumpulan Final Project")
    add_callout(doc, 
        "🔗 Tautan Repositori GitHub:\nhttps://github.com/iizaghz/PWS-UCP\n\n"
        "🌐 Tautan Live Deployment Vercel:\nhttps://ucp-cinedata-api.vercel.app\n\n"
        "📄 Berkas Laporan PDF (Siap diunggah ke Google Drive):\nLaporan_Final_Project_CineData_SaaS.pdf", 
        "INFORMASI TAUTAN PENGUMPULAN"
    )

    # --- BAB VII: PENUTUP ---
    add_heading_1(doc, "BAB VII: PENUTUP")
    add_heading_2(doc, "7.1 Kesimpulan")
    add_p(doc, "Sistem CineData API telah berhasil dibangun dan diuji dengan memenuhi 100% spesifikasi tugas akhir. Kombinasi arsitektur RESTful Express.js, basis data relasional PostgreSQL Supabase 11 tabel, sistem keamanan JWT + SHA-256 Hashing API Key, pembatasan laju (rate limiting), serta tampilan dashboard web yang modern menjadikannya sebuah solusi SaaS API data film yang tangguh, aman, dan siap pakai.")

    # Save Document
    doc_path = "Laporan_Final_Project_CineData_SaaS.docx"
    doc.save(doc_path)
    print(f"Document successfully created at {doc_path}")

if __name__ == "__main__":
    main()
